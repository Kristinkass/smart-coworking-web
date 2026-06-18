# -*- coding: utf-8 -*-
"""Add bibliography and in-text citations [N] to diploma PZ document."""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

DOCX_PATH = Path(
    r"c:\Users\user\Documents\4 КУРС\ДИПЛОМ))+))\ДИплом\Главы норм\ПЗ_ЦветковаК.А..docx"
)
BACKUP_PATH = DOCX_PATH.with_suffix(".backup.docx")
OUTPUT_PATH = DOCX_PATH.with_name("ПЗ_ЦветковаК.А._с_источниками_v4.docx")
ACCESS_DATE = "16.06.2026"

# Порядок: один ГОСТ → книги → англ. ресурсы → рус. ресурсы
BIBLIOGRAPHY = [
    "ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления : национальный стандарт Российской Федерации : издание официальное : утвержден и введен в действие Приказом Федерального агентства по техническому регулированию и метрологии от 27 апреля 2018 г. № 269-ст : введен впервые : дата введения 2019-07-01. – Москва : Стандартинформ, 2018. – 124 с. – Текст : непосредственный.",
    "Буч, Г. UML. Основы / Г. Буч, Дж. Рамбо, И. Джекобсон ; пер. с англ. – Санкт-Петербург : Питер, 2020. – 496 с. – Текст : непосредственный.",
    "Дейт, К. Дж. Введение в системы баз данных / К. Дж. Дейт ; пер. с англ. – 8-е изд. – Москва : Вильямс, 2020. – 1328 с. – Текст : непосредственный.",
    "Лутц, М. Изучаем Python / М. Лутц. – 5-е изд. – Санкт-Петербург : Вильямс, 2019. – 1280 с. – Текст : непосредственный.",
    "Мартин, Р. Чистая архитектура. Искусство разработки программного обеспечения / Р. Мартин. – Санкт-Петербург : Питер, 2023. – 352 с. – Текст : непосредственный.",
    "Соммервилл, И. Программная инженерия : учебник для вузов / И. Соммервилл. – Москва : ДМК Пресс, 2022. – 640 с. – Текст : непосредственный.",
    "Фаулер, М. Архитектура корпоративных программных приложений / М. Фаулер. – Москва : Вильямс, 2022. – 544 с. – Текст : непосредственный.",
    "Bootstrap 5 : документация : [Электронный ресурс]. – URL: https://getbootstrap.com/docs/5.3/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "BPMN 2.0 : спецификация нотации : [Электронный ресурс]. – URL: https://www.omg.org/spec/BPMN/2.0/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "Chart.js : документация : [Электронный ресурс]. – URL: https://www.chartjs.org/docs/latest/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "draw.io : инструмент визуального моделирования : [Электронный ресурс]. – URL: https://www.diagrams.net/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "Figma : облачный редактор интерфейсов : [Электронный ресурс]. – URL: https://www.figma.com/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "Flask : документация : [Электронный ресурс]. – URL: https://flask.palletsprojects.com/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "Global Coworking Space Statistics 2024 / Servcorp : [Электронный ресурс]. – URL: https://www.servcorp.com.au/en/blog/business-networking/how-many-coworking-spaces-are-there-in-the-world-2024-coworking-statistics/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "History of Coworking Spaces: from 2005 to 2021 / Coworking Resources : [Электронный ресурс]. – URL: https://www.coworkingresources.org/blog/history-of-coworking (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "Neuberg, B. Coworking — Community for Developers Who Work From Home : [Электронный ресурс]. – URL: http://codinginparadise.org/weblog/2005/08/coworking-community-for-developers-who.html (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "PostgreSQL 16 : документация : [Электронный ресурс]. – URL: https://www.postgresql.org/docs/16/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "PyCharm : интегрированная среда разработки : [Электронный ресурс]. – URL: https://www.jetbrains.com/pycharm/ (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
    "NF Group. Рынок гибких офисов. Москва. 2024 : [Электронный ресурс]. – URL: https://media.kf.expert/lenta_analytics/0/840/NF%20Group_%D0%A0%D1%8B%D0%BD%D0%BE%D0%BA%20%D0%B3%D0%B8%D0%B1%D0%BA%D0%B8%D1%85%20%D0%BE%D1%84%D0%B8%D1%81%D0%BE%D0%B2.%20%D0%9C%D0%BE%D1%81%D0%BA%D0%B2%D0%B0.%202024.pdf (дата обращения: "
    + ACCESS_DATE
    + "). – Текст : электронный.",
]

CITATION_RE = re.compile(r"\s*\[\d+(?:,\s*\d+)*\]")

CITATION_RULES: list[tuple[str, str]] = [
    ("Коворкинги стали эффективной альтернативой традиционным офисам", " [15]"),
    ("Одним из первых известных проектов считается сообщество C-base", " [15, 16]"),
    ("официально термин «коворкинг» получил распространение в 2005 году", " [15, 16]"),
    ("Если в 2018 году в мире насчитывалось около 16,6 тысячи коворкингов", " [14]"),
    ("к концу 2024 года число подобных пространств может приблизиться к 42 тысячам", " [14]"),
    ("Дополнительным фактором стала пандемия COVID-19", " [19]"),
    ("в 2024 году в Москве насчитывалось более 55 тысяч рабочих мест", " [19]"),
    ("Средняя заполняемость коворкингов достигла 95 % в Москве", " [19]"),
    ("стали внедряться специализированные информационные системы управления рабочими местами", " [6]"),
    ("Figma – это облачный графический редактор, выпущенный в 2016 году", " [12]"),
    ("Draw.io – бесплатный инструмент, который представляет собой универсальное решение", " [11]"),
    ("Проектирование системы является одним из ключевых этапов жизненного цикла разработки", " [6]"),
    ("было выполнено функциональное моделирование в нотации IDEF0", " [6]"),
    ("DFD-диаграмма – это визуальное представление того, как данные входят в систему", " [6]"),
    ("Диаграмма потоков данных (DFD) в нотации Гейна-Сарсона", " [6]"),
    ("была разработана диаграмма вариантов использования, представленная на рисунке 11", " [2]"),
    ("Диаграмма процесса создания бронирования построена в нотации BPMN", " [9]"),
    ("Система имеет трехуровневую клиент-серверную архитектуру", " [5, 7]"),
    ("Для представления о структуре базы данных была разработана диаграмма в нотации IDEF1X", " [3]"),
    ("Для хранения данных была выбрана система управления базами данных PostgreSQL", " [3, 17]"),
    ("Одним из основных преимуществ является полноценная поддержка ACID-транзакций", " [17]"),
    ("Разработка программного продукта осуществлялась на языке программирования Python", " [4]"),
    ("В качестве основной интегрированной среды разработки (IDE) использовался PyCharm", " [18]"),
    ("Для реализации серверной части применён фреймворк Flask", " [13]"),
    ("HTML5 и CSS3 обеспечивают разметку и оформление страниц, Bootstrap 5", " [8, 10]"),
    ("При проектировании серверной части применён многослойный подход", " [5, 7]"),
    ("Тестирование разработанного программного обеспечения является обязательным этапом", " [6]"),
    ("спроектировать веб-приложение на основе фреймворка Python Flask", " [4, 13]"),
    ("разработать модель структуры базы данных в нотации IDEF1X", " [3]"),
]


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text or ""


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def strip_citations(text: str) -> str:
    return CITATION_RE.sub("", text).rstrip()


def append_citation(paragraph: Paragraph, citation: str) -> bool:
    text = strip_citations(paragraph_text(paragraph))
    if citation.strip() in text:
        return False
    set_paragraph_text(paragraph, text + citation)
    return True


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def replace_bibliography(doc: Document) -> None:
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = paragraph_text(p).strip()
        if t == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ":
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError("Bibliography heading not found")

    to_remove = []
    for j in range(start_idx + 1, len(doc.paragraphs)):
        t = paragraph_text(doc.paragraphs[j]).strip()
        if t.startswith("ПРИЛОЖЕНИЕ"):
            break
        if t:
            to_remove.append(doc.paragraphs[j]._p)

    for element in to_remove:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    anchor = doc.paragraphs[start_idx]
    prev = anchor
    for entry in BIBLIOGRAPHY:
        prev = insert_paragraph_after(prev, entry, anchor.style)


def update_source_count_in_referat(doc: Document) -> None:
    count = len(BIBLIOGRAPHY)
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if "источник" in text and "090207" in text:
            new_text = re.sub(r"\d+\s+источник\w*", f"{count} источников", text)
            if new_text != text:
                set_paragraph_text(paragraph, new_text)
            break


def apply_citations(doc: Document) -> int:
    applied = 0
    for substring, citation in CITATION_RULES:
        for paragraph in doc.paragraphs:
            text = strip_citations(paragraph_text(paragraph))
            if substring in text:
                set_paragraph_text(paragraph, text)
                if append_citation(paragraph, citation):
                    applied += 1
                break
    return applied


def main() -> None:
    source = BACKUP_PATH if BACKUP_PATH.exists() else DOCX_PATH
    if not source.exists():
        raise FileNotFoundError(source)

    doc = Document(str(source))
    applied = apply_citations(doc)
    replace_bibliography(doc)
    update_source_count_in_referat(doc)
    doc.save(str(OUTPUT_PATH))

    print(f"Source: {source}")
    print(f"Citations applied: {applied}")
    print(f"Bibliography entries: {len(BIBLIOGRAPHY)}")
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
