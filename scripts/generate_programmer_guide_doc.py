# -*- coding: utf-8 -*-
"""Generate Programmer Guide (Руководство программиста) for coworking system."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT = "Система управления рабочими местами в коворкинге"
PRODUCT = "Умный Коворкинг"
AUTHOR = "Цветкова К.А."
GROUP = "ИСПк-403-52-00"
YEAR = "2026"


def _font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run(text), bold=True)


def para(doc, text, indent=True, bullet=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    prefix = f"{bullet} " if bullet else ""
    _font(p.add_run(prefix + text))
    return p


def bullets(doc, items):
    for item in items:
        para(doc, item, indent=False, bullet="–")


def sub_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        _font(p.add_run("– " + item))


def table2col(doc, title, rows, hdr=("Вид ресурса", "Минимальные требования")):
    para(doc, title, indent=False)
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = hdr[0]
    t.rows[0].cells[1].text = hdr[1]
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                _font(r, bold=True)
    for i, (a, b) in enumerate(rows, start=0):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    doc.add_paragraph()


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    _font(run, size=12)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")


def build() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(t.add_run("РУКОВОДСТВО ПРОГРАММИСТА"), bold=True)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(t2.add_run(PROJECT.lower()))
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(t3.add_run(YEAR))
    doc.add_paragraph()

    heading(doc, "АННОТАЦИЯ")
    para(
        doc,
        f"Настоящий документ содержит основные положения и сведения, необходимые программисту "
        f"для установки, настройки, сопровождения и развития веб-приложения «{PRODUCT}» — "
        f"{PROJECT.lower()}.",
    )
    para(
        doc,
        "Документ разработан согласно требованиям ГОСТ 19.504-79 "
        "«Руководство программиста. Требования к содержанию и оформлению».",
    )

    heading(doc, "Содержание")
    toc = [
        "1 Общие положения",
        "1.1 Наименование программы",
        "2 Назначение и условия применения",
        "2.1 Назначение программы",
        "2.2 Функции, выполняемые программой",
        "2.3 Сведения о технических и программных средствах",
        "2.3.1 Требования к составу и параметрам технических средств",
        "2.3.2 Требования к общесистемному программному обеспечению",
        "2.3.3 Настройка PostgreSQL и переменных окружения",
        "2.3.4 Требования к персоналу",
        "3 Характеристика программных средств программы",
        "3.1 Описание основных характеристик программы",
        "3.1.1 Временные характеристики программы",
        "3.1.2 Режим работы программы",
        "3.1.3 Средства контроля правильности выполнения программы",
        "4 Обращение к программе",
        "5 Входные и выходные данные",
        "5.1 Характер и организация входных данных",
        "5.2 Характер и организация выходных данных",
        "6 Сообщения",
        "7 Аварийные ситуации",
    ]
    for line in toc:
        para(doc, line, indent=False)
    doc.add_page_break()

    # 1
    heading(doc, "1 Общие положения")
    para(
        doc,
        f"Настоящий документ содержит сведения, необходимые для работы программиста "
        f"с {PROJECT.lower()}.",
    )
    heading(doc, "1.1 Наименование программы")
    para(doc, f"Полное наименование: {PROJECT}.")
    para(doc, f"Краткое наименование в интерфейсе: «{PRODUCT}».")
    para(doc, "Версия программы: 1.0.")

    # 2
    heading(doc, "2 Назначение и условия применения")
    heading(doc, "2.1 Назначение программы")
    para(
        doc,
        f"Основным назначением «{PRODUCT}» является автоматизация управления рабочими местами "
        f"в коворкинге, включая:",
    )
    bullets(
        doc,
        [
            "интерактивное бронирование через SVG-карту пространства;",
            "ведение клиентской базы с разграничением ролей (client, manager, admin);",
            "расчёт занятости по 15-минутной временной сетке и тарификацию;",
            "администрирование планировки, категорий мест, расписания и абонементов;",
            "формирование аналитической отчётности и экспорт в PDF.",
        ],
    )

    heading(doc, "2.2 Функции, выполняемые программой")
    para(doc, "Перечень функций, реализуемых программой:")
    bullets(
        doc,
        [
            "регистрация, аутентификация и авторизация пользователей (Flask-Login);",
            "отображение интерактивной карты с объединением layout.json и PostgreSQL;",
            "создание, продление и отмена бронирований с проверкой конфликтов;",
            "управление категориями, тарифами, расписанием и абонементами;",
            "визуальный редактор планировки с синхронизацией геометрии и БД;",
            "REST API для карты, бронирования, уведомлений и отчётов;",
            "генерация PDF-отчётов (ReportLab);",
            "автоматическая инициализация БД и начальных данных при первом запуске.",
        ],
    )

    heading(doc, "2.3 Сведения о технических и программных средствах")
    heading(doc, "2.3.1 Требования к составу и параметрам технических средств")
    table2col(
        doc,
        "Таблица 1 – Минимальные системные требования",
        [
            ("Тип ЭВМ", "Персональный компьютер или сервер"),
            ("Процессор", "Тактовая частота не менее 2.0 ГГц, 2 ядра"),
            ("Оперативная память (ОЗУ)", "Не менее 8 ГБ"),
            ("Накопитель (HDD/SSD)", "Не менее 2 ГБ свободного места"),
            ("Монитор", "Разрешение не ниже 1280×720"),
            ("Устройства ввода", "Клавиатура и мышь"),
            ("Операционная система", "Windows 10/11 или Linux"),
            ("Сетевое соединение", "Локальная сеть или localhost для доступа к веб-интерфейсу"),
        ],
    )

    heading(doc, "2.3.2 Требования к общесистемному программному обеспечению")
    table2col(
        doc,
        "Таблица 2 – Сведения о программном обеспечении",
        [
            ("Операционная система", "Windows 10/11 (64-разрядная) или Linux"),
            ("Интерпретатор Python", "Версия 3.13"),
            ("Веб-фреймворк", "Flask 2.3.3"),
            ("ORM", "Flask-SQLAlchemy 3.0.5"),
            ("Аутентификация", "Flask-Login 0.6.2"),
            ("СУБД", "PostgreSQL 16"),
            ("Драйвер БД", "psycopg2-binary ≥ 2.9.12"),
            ("Генерация PDF", "ReportLab 4.0.7"),
            ("Конфигурация", "python-dotenv 1.0.0"),
            ("Тестирование", "pytest 8.3.4"),
            ("Клиентская часть", "HTML5, CSS3, JavaScript, Bootstrap 5, SVG"),
        ],
    )

    para(
        doc,
        "Взаимодействие компонентов системы реализовано в рамках клиент–серверной архитектуры "
        "(рисунок 1). Серверная часть на Python/Flask обрабатывает HTTP-запросы и REST API; "
        "клиентская часть исполняется в браузере пользователя.",
    )
    para(doc, "Рисунок 1 – Архитектурная схема функционирования системы")
    para(doc, "В структуре программного комплекса выделены следующие компоненты:", indent=False)

    components = [
        (
            "Точка входа (cmd/app/main.py)",
            "запускает фабрику приложения create_app(), инициализирует БД (init_db) "
            "и встроенный веб-сервер Flask на порту 5000.",
        ),
        (
            "Слой application (internal/application.py)",
            "создаёт экземпляр Flask, подключает конфигурацию, SQLAlchemy, Flask-Login, "
            "регистрирует обработчики HTTP-запросов.",
        ),
        (
            "Слой handlers (internal/handlers/)",
            "принимает HTTP-запросы: web (страницы, auth), booking (API бронирования), "
            "map (данные карты), admin (панель управления), api (категории, отчёты, расписание).",
        ),
        (
            "Слой services (internal/services/)",
            "содержит бизнес-логику: booking_service (временные слоты, расчёт стоимости), "
            "room_editor_service (редактор планировки).",
        ),
        (
            "Слой repositories (internal/repositories/)",
            "инкапсулирует доступ к PostgreSQL: place_repository, booking_repository, user_repository.",
        ),
        (
            "Слой models (internal/models/)",
            "ORM-модели SQLAlchemy: User, Place, Booking, Subscription, Notification и др.",
        ),
        (
            "Слой layout (internal/layout/)",
            "чтение и кеширование static/layout.json, синхронизация геометрии с БД.",
        ),
        (
            "Клиентская часть (templates/, static/)",
            "HTML-шаблоны Jinja2, CSS, JavaScript-модули карты (map_updated.js), "
            "бронирования (booking.js) и редактора (editor/).",
        ),
    ]
    for name, desc in components:
        para(doc, f"{name} — {desc}")

    heading(doc, "2.3.3 Настройка PostgreSQL и переменных окружения")
    para(doc, "Порядок подготовки серверной среды:", indent=False)
    bullets(
        doc,
        [
            "установить PostgreSQL 16 и создать пользователя с правами CREATE DATABASE;",
            "скопировать .env.example в .env в корне проекта;",
            "заполнить параметры DB_HOST, DB_USER, DB_PASSWORD, DB_NAME, DB_PORT;",
            "указать SECRET_KEY и учётные данные первого администратора (ADMIN_EMAIL, ADMIN_PASSWORD);",
            "выполнить: pip install -r requirements.txt;",
            "запустить: py cmd/app/main.py.",
        ],
    )
    para(doc, "Пример содержимого файла .env:", indent=False)
    code_block(
        doc,
        "SECRET_KEY=ваша-случайная-строка\n"
        "DB_HOST=localhost\nDB_USER=postgres\nDB_PASSWORD=пароль\n"
        "DB_NAME=coworking\nDB_PORT=5432\n"
        "ADMIN_EMAIL=admin@example.com\nADMIN_PASSWORD=ChangeMe123!",
    )
    para(
        doc,
        "При каждом запуске (кроме режима TESTING) функция init_db() создаёт отсутствующие таблицы "
        "(db.create_all()), выполняет ad-hoc-миграции схемы (run_migrations) для обновления "
        "устаревших баз данных, загружает начальные данные (init_default_data) при их отсутствии "
        "и обновляет статусы бронирований. Функция ensure_database() создаёт базу PostgreSQL, "
        "если она ещё не существует.",
    )
    para(doc, "Проверка подключения к PostgreSQL (Листинг 1):", indent=False)
    code_block(
        doc,
        "import psycopg2\n"
        "conn = psycopg2.connect(host='localhost', user='postgres',\n"
        "                        password='пароль', dbname='coworking')\n"
        "print('Подключение успешно')\nconn.close()",
    )

    heading(doc, "2.3.4 Требования к персоналу")
    para(doc, "Обязательные системные навыки программиста:", indent=False)
    bullets(
        doc,
        [
            "навыки работы в ОС Windows/Linux (развёртывание окружения, виртуальные среды Python venv);",
            "уверенное владение терминалом для установки зависимостей и запуска приложения;",
            "базовые навыки администрирования PostgreSQL.",
        ],
    )
    para(doc, "Специальные профессиональные навыки:", indent=False)
    bullets(
        doc,
        [
            "Python-разработка: знание Flask, паттерна application factory, Blueprint;",
            "ORM SQLAlchemy: проектирование моделей, миграции, транзакции;",
            "SQL: составление запросов SELECT/INSERT/UPDATE/DELETE, понимание FK и индексов;",
            "веб-разработка: HTML, CSS, JavaScript, REST API, JSON;",
            "понимание клиент–серверной архитектуры и разграничения доступа по ролям;",
            "опыт работы с pytest для регрессионного тестирования.",
        ],
    )

    # 3
    heading(doc, "3 Характеристика программных средств программы")
    heading(doc, "3.1 Описание основных характеристик программы")
    para(
        doc,
        f"Программное обеспечение «{PRODUCT}» реализовано как веб-приложение с многослойной "
        f"серверной архитектурой. Данные разделены на две части: транзакционные сущности "
        f"(пользователи, бронирования, тарифы) хранятся в PostgreSQL; геометрия планировки "
        f"(координаты мест, стены, двери) — в файле static/layout.json. Связь осуществляется по коду места.",
    )
    para(doc, "Структура каталогов проекта:", indent=False)
    bullets(
        doc,
        [
            "cmd/app/main.py — точка запуска;",
            "internal/application.py — фабрика Flask-приложения;",
            "internal/config.py — конфигурация и TestConfig;",
            "internal/handlers/ — HTTP-обработчики по подсистемам;",
            "internal/services/ — бизнес-логика;",
            "internal/repositories/ — доступ к данным;",
            "internal/models/ — ORM-модели и seed;",
            "internal/layout/ — работа с layout.json;",
            "internal/utils/ — декораторы, форматтеры, утилиты;",
            "templates/ — HTML-шаблоны;",
            "static/js/ — клиентские модули;",
            "static/layout.json — геометрия планировки;",
            "tests/ — автоматизированные тесты pytest.",
        ],
    )

    heading(doc, "3.1.1 Временные характеристики программы")
    bullets(
        doc,
        [
            "среднее время запуска сервера — не более 30 секунд (включая инициализацию БД);",
            "время ответа REST API /api/places — менее 500 мс при типичной планировке;",
            "время отклика веб-интерфейса при типовых операциях — не более 2 секунд;",
            "минимальная длительность бронирования — 30 минут (2 слота по 15 минут);",
            "регрессионный прогон pytest (tests/) — не более 120 секунд.",
        ],
    )

    heading(doc, "3.1.2 Режим работы программы")
    para(
        doc,
        "Программа функционирует в интерактивном режиме через веб-браузер. "
        "Сервер обрабатывает синхронные HTTP-запросы; динамические операции карты и бронирования "
        "выполняются асинхронно на клиенте через fetch/XMLHttpRequest к REST API. "
        "При DEBUG=True Flask запускается с одним рабочим процессом; для production "
        "рекомендуется развёртывание за WSGI-сервером (Gunicorn, Waitress).",
    )

    heading(doc, "3.1.3 Средства контроля правильности выполнения программы")
    bullets(
        doc,
        [
            "валидация входных данных в обработчиках и сервисном слое (формат даты, email, длительность брони);",
            "декораторы @admin_required, @staff_required, @manager_required для RBAC;",
            "проверка конфликтов бронирования в booking_service перед записью в БД;",
            "транзакции SQLAlchemy с rollback при ошибках;",
            "автоматизированные тесты pytest в каталоге tests/ (58+ тестов);",
            "flash-сообщения и JSON-ответы с полем error для REST API;",
            "логирование ошибок в консоль сервера при DEBUG=True.",
        ],
    )

    # 4
    heading(doc, "4 Обращение к программе")
    para(
        doc,
        f"Программа «{PRODUCT}» запускается вызовом точки входа cmd/app/main.py "
        f"через интерпретатор Python.",
    )
    para(doc, "Порядок вызова программы:", indent=False)
    bullets(
        doc,
        [
            "открыть терминал в корневом каталоге проекта;",
            "убедиться, что файл .env заполнен и PostgreSQL доступен;",
            "выполнить: py cmd/app/main.py;",
            "дождаться сообщения «Running on http://127.0.0.1:5000»;",
            "открыть указанный адрес в браузере.",
        ],
    )
    para(doc, "Последовательность инициализации при запуске:", indent=False)
    bullets(
        doc,
        [
            "загрузка переменных окружения из .env (python-dotenv);",
            "создание экземпляра Flask через create_app();",
            "подключение SQLAlchemy к PostgreSQL;",
            "инициализация Flask-Login и регистрация маршрутов (register_all_handlers);",
            "выполнение init_db(): создание таблиц, run_migrations(), seed, обновление статусов броней;",
            "запуск встроенного веб-сервера Werkzeug на порту 5000.",
        ],
    )
    para(doc, "Запуск в production (Windows — Waitress, Linux — Gunicorn):", indent=False)
    code_block(doc, "waitress-serve --host=0.0.0.0 --port=5000 wsgi:app")
    para(doc, "Запуск автоматизированных тестов:", indent=False)
    code_block(doc, "py -m pytest tests/ -q")

    # 5
    heading(doc, "5 Входные и выходные данные")
    heading(doc, "5.1 Характер и организация входных данных")
    para(doc, "Входными данными являются:", indent=False)
    bullets(
        doc,
        [
            "HTTP-запросы пользователя (GET/POST) с параметрами форм и JSON-телами REST API;",
            "учётные данные: email или телефон, пароль (хешируется Werkzeug);",
            "параметры бронирования: place_id, дата, время начала/окончания, тип тарифа, число посетителей;",
            "геометрические данные редактора: координаты, размеры, коды мест в layout.json;",
            "конфигурация окружения: переменные .env (DB_*, SECRET_KEY, ADMIN_*).",
        ],
    )
    para(
        doc,
        "Основные таблицы PostgreSQL (входное/хранимое состояние): users, places, bookings, "
        "place_categories, category_tariffs, subscriptions, notifications, coworking_schedules. "
        "Полное описание схемы — в документе Database-schema.md.",
    )

    heading(doc, "5.2 Характер и организация выходных данных")
    para(
        doc,
        "Выходными данными являются результаты работы подсистем бронирования, администрирования "
        "и аналитики, формируемые в виде веб-страниц, JSON-структур, файлов и записей СУБД.",
    )
    para(doc, "Основные виды выходных данных:", indent=False)

    para(doc, "HTML-страницы (шаблоны Jinja2, каталог templates/):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "содержат сформированный интерфейс главной страницы, личного кабинета, карты бронирования "
            "(/mapp), панели администратора и редактора планировки;",
            "передаются браузеру с кодом HTTP 200 при успешном рендеринге; при ошибках доступа "
            "возвращаются редиректы (302) или страницы с flash-сообщениями;",
            "включают встроенные блоки уведомлений (flash-container в base.html) и данные "
            "для клиентских скриптов (карта, бронирование, профиль).",
        ],
    )

    para(doc, "JSON-ответы REST API (Content-Type: application/json):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "/api/places — геометрия и статусы мест: массивы places, walls, doors; каждое место "
            "содержит code, x, y, width, height, status, category, bookings, location;",
            "/api/booking/timegrid/<place_id> — 15-минутная сетка занятости на выбранную дату "
            "с признаками available, price, conflict;",
            "/api/booking/create, /api/booking/price — результат расчёта и создания бронирования "
            "(success, total_price, booking_id или error);",
            "/api/notifications — списки системных уведомлений и обращений клиентов (kind: system/feedback);",
            "/api/admin/reports — агрегированная статистика бронирований за период (JSON перед экспортом);",
            "при ошибках API возвращает {\"success\": false, \"error\": \"текст\"} с кодами 400, 401, 403, 404, 500.",
        ],
    )

    para(doc, "Документы PDF (ReportLab, маршрут /api/admin/reports/pdf):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "готовые к печати отчёты по бронированиям за выбранный период (параметры start_date, end_date, type);",
            "содержат заголовок отчёта, сводную статистику (число броней, сумма, средняя длительность), "
            "таблицы по разделам: почасовые, абонементные, завершённые, активные, отменённые;",
            "каждая строка таблицы включает дату, время/период, пользователя, место, локацию, "
            "длительность, сумму и статус бронирования;",
            "файл передаётся браузеру как вложение (Content-Disposition: attachment, формат A4).",
        ],
    )

    para(doc, "Файл геометрии планировки (static/layout.json):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "специализированный JSON-формат хранения координат мест, стен, дверей и метаданных этажей;",
            "обновляется при сохранении изменений в визуальном редакторе (internal/layout/store.py);",
            "синхронизируется с таблицами PostgreSQL (places, locations) по коду места (code);",
            "резервная копия сохраняется в static/layout.json.backup перед критическими операциями.",
        ],
    )

    para(doc, "Записи в PostgreSQL (транзакционное состояние системы):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "бронирования (bookings): дата, интервал, стоимость, статус (active/completed/cancelled);",
            "пользователи (users): учётные данные, роль, статус active;",
            "уведомления (notifications): массовые и персональные сообщения, обращения клиентов;",
            "абонементы (subscriptions), тарифы (category_tariffs), расписание (coworking_schedules).",
        ],
    )

    para(doc, "HTTP-ответы с кодами состояния:", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "200 — успешное выполнение запроса;",
            "201 — успешное создание ресурса (бронирование, абонемент);",
            "302 — перенаправление после POST (вход, выход, админ-действия);",
            "400/401/403/404/500 — ошибки валидации, авторизации, доступа, отсутствия ресурса, сбоя сервера.",
        ],
    )

    para(doc, "Консольный вывод сервера (stdout при запуске):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "статус инициализации БД: «Создание/проверка таблиц…», «[OK] Коворкинг», «[OK] Рабочие места»;",
            "сообщения ad-hoc-миграций: «[MIGRATE] …»;",
            "предупреждения синхронизации layout: «[ERR] Локация не найдена: …»;",
            "сообщение Waitress/Gunicorn: «Serving on http://0.0.0.0:5000».",
        ],
    )

    # 6
    heading(doc, "6 Сообщения")
    para(
        doc,
        "Сообщения в системе реализуются средствами веб-интерфейса (HTML-шаблоны Jinja2, CSS, "
        "JavaScript-модуль static/js/notify.js, механизм Flask flash) и делятся на следующие "
        "категории в зависимости от критичности событий:",
    )

    para(doc, "Информационные сообщения (категория success):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "реализуются через flash('…', 'success') после POST-запросов и через showToast(…, 'success') "
            "в клиентских скриптах;",
            "отображаются в правом верхнем углу экрана (flash-container, cw-toast) с зелёной "
            "цветовой индикацией и иконкой fa-check-circle;",
            "уведомляют пользователя об успешном завершении операций;",
            "примеры: «Добро пожаловать, …!», «Регистрация успешна!», «Бронирование успешно отменено», "
            "«Пароль успешно изменён», «Бронирование создано» (toast на карте).",
        ],
    )

    para(doc, "Предупреждения (категория warning):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "реализуются через flash('…', 'warning') и showToast(…, 'warning');",
            "сообщают о некритичных проблемах, не блокирующих работу программы полностью;",
            "отображаются с жёлтой индикацией и иконкой fa-exclamation-triangle;",
            "примеры: «Дата окончания скорректирована до сегодняшнего дня», «Пароли не совпадают», "
            "«Введите корректный номер телефона», «Имя и телефон обязательны».",
        ],
    )

    para(doc, "Ошибки (категория error):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "реализуются через flash('…', 'error'), showToast(…, 'error') и JSON-ответы REST API "
            "с полем error;",
            "информируют о невозможности продолжения текущей операции из-за неверных действий "
            "пользователя, нарушения прав доступа или сбоев среды;",
            "отображаются с красной индикацией и иконкой fa-exclamation-circle / fa-times-circle;",
            "примеры: «Доступ запрещен. Требуются права администратора.», "
            "«Неверная почта или пароль. Попробуйте снова.», «Ваш аккаунт деактивирован», "
            "«Минимальная длительность бронирования — 30 минут», «Выбранное время недоступно», "
            "«Ошибка при загрузке карты: …».",
        ],
    )

    para(doc, "Справочные сообщения (категория info):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "реализуются через flash('…', 'info') и showToast(…, 'info');",
            "используются для нейтральных подсказок в редакторе планировки и на формах;",
            "отображаются с синей индикацией и иконкой fa-info-circle;",
            "примеры: «Начало стены – кликните конец (можно по другой стене)» (редактор), "
            "информационные блоки на страницах администрирования.",
        ],
    )

    para(doc, "Диалоги подтверждения (showConfirm в notify.js):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "модальные окна с кнопками «Да» / «Отмена» для действий, требующих явного согласия пользователя;",
            "применяются при отмене бронирования, удалении обращений, критических операциях редактора;",
            "при отказе операция не выполняется; при подтверждении — продолжается с соответствующим API-запросом.",
        ],
    )

    para(doc, "Системные уведомления в базе данных (модель Notification):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "массовые рассылки администрации (target_audience: all, clients, managers, admins);",
            "персональные сообщения конкретному пользователю (user_id);",
            "обращения клиентов менеджерам и администраторам (kind: feedback, API /api/feedback);",
            "отображаются в личном кабинете (dashboard) и через GET /api/notifications с пометкой is_read.",
        ],
    )

    para(doc, "Консольные логи (вывод в stdout терминала):", indent=False, bullet="–")
    sub_bullets(
        doc,
        [
            "предназначены для программиста при отладке и сопровождении исходного кода;",
            "выводятся в консоль запуска сервера и содержат техническую информацию: ход инициализации "
            "PostgreSQL и seed-данных, сообщения ad-hoc-миграций ([MIGRATE]), предупреждения синхронизации "
            "layout ([ERR]), автоматическое завершение просроченных бронирований;",
            "при DEBUG=True Flask/Werkzeug дополнительно выводит трассировки исключений Python (traceback) "
            "и журнал HTTP-запросов (метод, URL, код ответа).",
        ],
    )

    # 7
    heading(doc, "7 Аварийные ситуации")
    para(
        doc,
        "Ниже перечислены типовые аварийные ситуации, возможные сообщения об ошибках "
        "и рекомендуемые действия программиста по их устранению.",
    )

    heading(doc, "7.1 Действия при ошибках подключения к PostgreSQL")
    para(doc, "Возможные ошибки и симптомы:", indent=False)
    sub_bullets(
        doc,
        [
            "psycopg2.OperationalError: connection refused — служба PostgreSQL не запущена "
            "или указан неверный DB_PORT;",
            "psycopg2.OperationalError: password authentication failed — неверный DB_USER или DB_PASSWORD в .env;",
            "psycopg2.OperationalError: database \"…\" does not exist — база DB_NAME отсутствует "
            "и автоматическое создание (ensure_database) не выполнилось;",
            "psycopg2.OperationalError: could not connect to server — неверный DB_HOST "
            "(localhost вместо IP сервера или наоборот);",
            "sqlalchemy.exc.OperationalError при HTTP-запросе — потеря соединения с БД "
            "после простоя (pool_pre_ping должен переподключать, но служба могла быть остановлена);",
            "UnicodeDecodeError / кодировка — некорректная локаль клиента "
            "(решается PGCLIENTENCODING=UTF8 в точке входа);",
            "ошибка при init_db(): «[MIGRATE] Ошибка: …» — сбой ad-hoc-миграции схемы "
            "на повреждённой или несовместимой базе.",
        ],
    )
    para(doc, "Рекомендуемые действия:", indent=False)
    bullets(
        doc,
        [
            "проверить, что служба PostgreSQL запущена;",
            "убедиться в корректности DB_HOST, DB_PORT, DB_USER, DB_PASSWORD в .env;",
            "проверить существование базы DB_NAME или разрешить автоматическое создание (ensure_database);",
            "просмотреть трассировку psycopg2.OperationalError в консоли сервера;",
            "проверить подключение отдельным скриптом (Листинг 1, раздел 2.3.3);",
            "при ошибках миграции — восстановить БД из резервной копии или выполнить pg_dump перед правкой схемы.",
        ],
    )

    heading(doc, "7.2 Действия при ошибках загрузки layout.json")
    para(doc, "Возможные ошибки и симптомы:", indent=False)
    sub_bullets(
        doc,
        [
            "FileNotFoundError: static/layout.json — файл планировки отсутствует или путь LAYOUT_PATH неверен;",
            "json.JSONDecodeError — повреждённый JSON (лишняя запятая, незакрытая скобка, неверная кодировка);",
            "KeyError при обращении к places, walls, floors — неполная структура layout "
            "(отсутствуют обязательные ключи);",
            "«[ERR] Локация не найдена: … (место …)» при seed — место в layout.json ссылается "
            "на несуществующую локацию в БД;",
            "пустая или некорректная карта в /mapp — кеш _LAYOUT_CACHE содержит устаревшие данные "
            "после ручного редактирования файла;",
            "HTTP 500 на /api/places или /mapp с текстом «Ошибка при загрузке карты» — "
            "исключение в internal/handlers/map/places.py при merge layout и PostgreSQL;",
            "несоответствие кодов мест (code) между layout.json и таблицей places — "
            "места не отображаются или не бронируются.",
        ],
    )
    para(doc, "Рекомендуемые действия:", indent=False)
    bullets(
        doc,
        [
            "проверить наличие файла static/layout.json;",
            "убедиться в корректности JSON-синтаксиса (валидатор JSON или json.load в Python);",
            "восстановить из резервной копии static/layout.json.backup;",
            "перезапустить сервер для сброса кеша (reload_layout() в internal/layout/store.py);",
            "при [ERR] Локация не найдена — создать локацию в админ-панели или исправить поле location в layout;",
            "синхронизировать места через редактор планировки или init_default_data после восстановления файла.",
        ],
    )

    heading(doc, "7.3 Действия при сбое веб-сервера")
    para(doc, "Возможные ошибки и симптомы:", indent=False)
    sub_bullets(
        doc,
        [
            "OSError: [WinError 10048] / Address already in use — порт 5000 занят другим процессом "
            "(Flask, Waitress, другое приложение);",
            "ModuleNotFoundError: No module named 'fcntl' — попытка запуска Gunicorn на Windows "
            "(Gunicorn поддерживает только Linux/macOS);",
            "ImportError / ModuleNotFoundError при старте — не установлены зависимости "
            "(pip install -r requirements.txt) или неверный каталог запуска;",
            "waitress-serve / py не распознаны — виртуальное окружение не активировано "
            "или Waitress не установлен;",
            "браузер: «Не удаётся подключиться» при Serving on http://0.0.0.0:5000 — "
            "открыт адрес http://0.0.0.0:5000 вместо http://127.0.0.1:5000 или http://<IP>:5000;",
            "доступ с телефона по LAN не работает — брандмауэр Windows блокирует входящие на порт 5000;",
            "HTTP 500 Internal Server Error — необработанное исключение Python (traceback в консоли);",
            "процесс Python завершился без сообщения — критическая ошибка при init_db() до старта сервера.",
        ],
    )
    para(doc, "Рекомендуемые действия:", indent=False)
    bullets(
        doc,
        [
            "остановить процесс Python (Ctrl+C в терминале или диспетчер задач);",
            "проверить, что порт 5000 не занят: netstat -ano | findstr :5000 (Windows);",
            "на Windows использовать Waitress: waitress-serve --host=0.0.0.0 --port=5000 wsgi:app;",
            "для разработки перезапустить: py cmd/app/main.py;",
            "открывать http://127.0.0.1:5000 локально или http://<IP-ПК>:5000 в LAN;",
            "разрешить приложению в брандмауэре Windows доступ к частной сети;",
            "при повторении ошибки — просмотреть traceback в консоли и исправить указанный модуль.",
        ],
    )

    heading(doc, "7.4 Действия при потере данных")
    para(doc, "Возможные ошибки и симптомы:", indent=False)
    sub_bullets(
        doc,
        [
            "случайное удаление или перезапись static/layout.json — карта пуста или искажена;",
            "повреждение таблиц PostgreSQL после сбоя диска или некорректной миграции;",
            "IntegrityError: violates foreign key constraint — попытка удалить место/пользователя "
            "с активными бронированиями (ON DELETE RESTRICT);",
            "отсутствие истории бронирований после восстановления — восстановлена устаревшая "
            "резервная копия без таблицы bookings;",
            "расхождение layout.json и БД — места на карте не совпадают с записями places;",
            "потеря уведомлений или абонементов — восстановлен partial dump без таблиц "
            "notifications / subscriptions.",
        ],
    )
    para(doc, "Рекомендуемые действия:", indent=False)
    bullets(
        doc,
        [
            "восстановить PostgreSQL из резервной копии (pg_dump / pg_restore);",
            "восстановить static/layout.json из static/layout.json.backup;",
            "пользователей в интерфейсе не удаляют — только блокируют (users.active = FALSE);",
            "история бронирований защищена ограничениями ON DELETE RESTRICT — "
            "перед удалением места отменить или завершить связанные брони;",
            "после восстановления выполнить init_db() и проверить синхронизацию layout ↔ places;",
            "настроить регулярное резервное копирование: pg_dump -Fc coworking > backup.dump.",
        ],
    )

    heading(doc, "7.5 Действия в других аварийных ситуациях")
    para(doc, "Возможные ошибки и симптомы:", indent=False)
    sub_bullets(
        doc,
        [
            "JSON {\"success\": false, \"error\": \"…\"} с кодом 400 — ошибка валидации входных данных "
            "(дата, время, place_id, email, телефон);",
            "HTTP 401 Unauthorized — сессия истекла, пользователь не выполнил вход (Flask-Login);",
            "HTTP 403 Forbidden — недостаточно прав (декораторы @admin_required, @staff_required);",
            "flash «Доступ запрещен…» после перехода в админ-раздел под ролью client;",
            "ошибки бронирования: «Минимальная длительность — 30 минут», «Выбранное время недоступно», "
            "«Отмена доступна не позднее чем за 1 час» — бизнес-правила booking_service;",
            "ошибка генерации PDF (ReportLab): отсутствует шрифт или некорректный период отчёта;",
            "падение pytest — регрессия после изменения handler/service/repository;",
            "404 на /static/js/jquery.min.js — отсутствует локальный файл (не критично, если не используется).",
        ],
    )
    para(doc, "Рекомендуемые действия:", indent=False)
    bullets(
        doc,
        [
            "определить причину сбоя по логам сервера, flash-сообщению и телу JSON-ответа API;",
            "выполнить py -m pytest tests/ -q для проверки регрессий;",
            "устранить дефект в соответствующем слое (handler → service → repository);",
            "повторить сценарий из ПМИ (раздел 3.6) после исправления;",
            "зафиксировать изменения в системе контроля версий и обновить документацию при изменении API.",
        ],
    )

    return doc


def main():
    out = Path(__file__).resolve().parents[1] / "docs"
    out.mkdir(exist_ok=True)
    path = out / "Руководство_программиста_ЦветковаК.А._система_управления_рабочими_местами.docx"
    alt = out / "Руководство_программиста_ЦветковаК.А._система_управления_рабочими_местами_updated.docx"
    doc = build()
    try:
        doc.save(path)
        print(f"Saved: {path}")
    except PermissionError:
        doc.save(alt)
        print(f"Saved (original file locked): {alt}")


if __name__ == "__main__":
    main()
